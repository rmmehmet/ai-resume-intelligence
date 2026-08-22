"""
Resume layout / ATS-parsability analysis.

Real ATS parsing engines (the kind behind Workday, Greenhouse,
iCIMS, Taleo, etc.) extract text in reading order and are well
known to fail on specific structural patterns:

- Multi-column layouts: text from two columns often gets read
  left-to-right across the page instead of down each column,
  scrambling the content.
- Tables: table cell contents are frequently dropped, merged, or
  reordered.
- Contact info placed only in a document header/footer: many ATS
  parsers strip headers/footers before parsing, silently discarding
  a candidate's name/email/phone.
- Image-based content (a resume that's mostly a picture/graphic):
  there's no text to extract at all.

These checks operate on the file's actual layout (positions of
words on the page / document XML structure), not just its text -
that's what makes them meaningfully different from the content-only
rules in services/ats/rules.py.
"""
import re
from io import BytesIO

from docx import Document

from schemas.layout import LayoutAnalysis

_EMAIL_PATTERN = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PHONE_PATTERN = re.compile(r"(\+?\d[\d\s().-]{7,}\d)")

# Header/footer zones as a fraction of page height, measured from the
# respective edge. Chosen generously to match typical resume margins.
_HEADER_ZONE_FRACTION = 0.12
_FOOTER_ZONE_FRACTION = 0.08

# A gap wider than normal word-spacing (typically 2-6pt) but appearing
# at a consistent x-position across many lines is what distinguishes a
# real column gutter from an incidental wide space. The absolute
# threshold matters more than a page-width fraction here, since
# word-spacing scales with font size, not page width.
_MIN_GUTTER_GAP_PT = 10.0
_COLUMN_MIN_SIDE_FRACTION = 0.15


def analyze_layout(content: bytes, extension: str) -> LayoutAnalysis:
    if extension == "pdf":
        return _analyze_pdf_layout(content)
    if extension == "docx":
        return _analyze_docx_layout(content)
    return LayoutAnalysis(notes=[f"No layout analysis available for .{extension}"])


def _garbled_ratio(text: str) -> float:
    """Proportion of characters that suggest bad encoding/extraction (e.g. an image-based PDF)."""
    if not text:
        return 1.0
    bad = sum(1 for ch in text if ch == "\ufffd" or (ord(ch) < 32 and ch not in "\n\t\r"))
    return round(bad / len(text), 4)


def _contains_contact(text: str) -> bool:
    return bool(_EMAIL_PATTERN.search(text) or _PHONE_PATTERN.search(text))


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------

def _analyze_pdf_layout(content: bytes) -> LayoutAnalysis:
    import pdfplumber

    notes: list[str] = []
    multi_column = False
    has_tables = False
    header_footer_only_contact = False
    all_text_parts: list[str] = []

    with pdfplumber.open(BytesIO(content)) as pdf:
        header_zone_text: list[str] = []
        footer_zone_text: list[str] = []
        body_zone_text: list[str] = []

        for page_index, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            all_text_parts.append(page_text)

            if page.find_tables():
                has_tables = True

            if _detect_multi_column(page):
                multi_column = True
                notes.append(f"Possible multi-column layout detected on page {page_index + 1}")

            header_h = page.height * _HEADER_ZONE_FRACTION
            footer_h = page.height * _FOOTER_ZONE_FRACTION
            for word in page.extract_words():
                if word["top"] <= header_h:
                    header_zone_text.append(word["text"])
                elif word["bottom"] >= page.height - footer_h:
                    footer_zone_text.append(word["text"])
                else:
                    body_zone_text.append(word["text"])

        header_footer_text = " ".join(header_zone_text + footer_zone_text)
        body_text = " ".join(body_zone_text)
        if _contains_contact(header_footer_text) and not _contains_contact(body_text):
            header_footer_only_contact = True
            notes.append("Contact info appears only in the header/footer zone")

    if has_tables:
        notes.append("Table structure detected - table content is often dropped or reordered by ATS parsers")

    full_text = "\n".join(all_text_parts)
    garbled = _garbled_ratio(full_text)
    if garbled > 0.02 or len(full_text.strip()) < 50:
        notes.append("Little or no extractable text found - resume may be image-based")

    return LayoutAnalysis(
        multi_column=multi_column,
        has_tables=has_tables,
        contact_only_in_header_footer=header_footer_only_contact,
        garbled_text_ratio=garbled,
        notes=notes,
    )


def _detect_multi_column(page) -> bool:
    """
    Heuristic: for each line of text (words sharing a vertical position),
    look for a gap between two words wider than normal word-spacing - a
    candidate "gutter" separating two columns. If a consistent fraction
    of lines share a gutter at roughly the same x-position, that's
    strong evidence of a genuine two-column layout rather than
    incidental spacing (e.g. a date right-aligned next to a job title,
    which would only affect a line or two, not a consistent band).
    """
    words = page.extract_words()
    if len(words) < 20:
        return False

    lines: dict[int, list[dict]] = {}
    for word in words:
        # Round `top` to bucket words into the same line despite tiny
        # sub-pixel differences.
        key = round(word["top"] / 3)
        lines.setdefault(key, []).append(word)

    gutter_midpoints: list[float] = []
    for line_words in lines.values():
        if len(line_words) < 2:
            continue
        ordered = sorted(line_words, key=lambda w: w["x0"])
        for a, b in zip(ordered, ordered[1:]):
            gap = b["x0"] - a["x1"]
            if gap >= _MIN_GUTTER_GAP_PT:
                gutter_midpoints.append((a["x1"] + b["x0"]) / 2)

    min_lines_with_gutter = max(3, int(len(lines) * _COLUMN_MIN_SIDE_FRACTION))
    if len(gutter_midpoints) < min_lines_with_gutter:
        return False

    # Gutters from a real column layout cluster tightly around one
    # x-position; scattered incidental gaps (e.g. a right-aligned date)
    # won't. Check that enough detected gutters fall within a narrow band.
    gutter_midpoints.sort()
    band_width = page.width * 0.05
    best_cluster_size = 1
    for base in gutter_midpoints:
        cluster_size = sum(1 for x in gutter_midpoints if base <= x <= base + band_width)
        best_cluster_size = max(best_cluster_size, cluster_size)

    return best_cluster_size >= min_lines_with_gutter


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------

def _analyze_docx_layout(content: bytes) -> LayoutAnalysis:
    document = Document(BytesIO(content))
    notes: list[str] = []

    has_tables = len(document.tables) > 0
    if has_tables:
        notes.append("Table structure detected - table content is often dropped or reordered by ATS parsers")

    multi_column = _docx_has_multiple_columns(document)
    if multi_column:
        notes.append("Multi-column section layout detected")

    header_text = "\n".join(
        p.text for section in document.sections for p in section.header.paragraphs
    )
    footer_text = "\n".join(
        p.text for section in document.sections for p in section.footer.paragraphs
    )
    body_text = "\n".join(p.text for p in document.paragraphs)

    header_footer_only_contact = False
    if _contains_contact(header_text + " " + footer_text) and not _contains_contact(body_text):
        header_footer_only_contact = True
        notes.append("Contact info appears only in the document header/footer")

    garbled = _garbled_ratio(body_text)
    if len(body_text.strip()) < 50:
        notes.append("Little or no extractable text found in the document body")

    return LayoutAnalysis(
        multi_column=multi_column,
        has_tables=has_tables,
        contact_only_in_header_footer=header_footer_only_contact,
        garbled_text_ratio=garbled,
        notes=notes,
    )


def _docx_has_multiple_columns(document: Document) -> bool:
    """A DOCX section can define a column count via its sectPr/cols XML element."""
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for section in document.sections:
        for cols in section._sectPr.findall("w:cols", ns):
            num = cols.get(f"{{{ns['w']}}}num")
            if num and int(num) > 1:
                return True
    return False